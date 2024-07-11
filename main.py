import streamlit as st
from io import BytesIO
import base64
import requests

import openai

import cv2
import numpy as np
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from PIL import Image
#import pytesseract

import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.datavalidation import DataValidation
import pandas as pd
from openpyxl import Workbook
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.styles import PatternFill, Font, Alignment


from docx import Document
from docx.shared import Inches

import io

import torch
import torchaudio
from transformers import AutoProcessor, AutoModelForSpeechSeq2Seq


def transcribe_czech_audio(audio_file_path, segment_length=30):
    """
    Finetuned model for transcribing Czech audio
    """
    # Load the processor and the model
    processor = AutoProcessor.from_pretrained("Cem13/whisper-large-v3-czech")
    model = AutoModelForSpeechSeq2Seq.from_pretrained("Cem13/whisper-large-v3-czech")

    # Load and preprocess the audio file
    try:
        waveform, sample_rate = torchaudio.load(audio_file_path)
        print(f"Loaded audio file: {audio_file_path}")
    except FileNotFoundError:
        print(f"Error: Audio file '{audio_file_path}' not found.")
        return None

    # Convert to mono if necessary
    if waveform.shape[0] > 1:
        waveform = torch.mean(waveform, dim=0, keepdim=True)
        print("Converted audio to mono")

    # Resample if sample rate doesn't match model's expectation
    expected_sample_rate = processor.feature_extractor.sampling_rate
    if sample_rate != expected_sample_rate:
        waveform = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=expected_sample_rate)(waveform)
        print(f"Resampled audio from {sample_rate}Hz to {expected_sample_rate}Hz")

    # Calculate the number of segments
    total_length = waveform.shape[1] / expected_sample_rate
    num_segments = int(total_length // segment_length) + 1
    print(f"Total length of audio: {total_length:.2f}s, Number of segments: {num_segments}")

    transcriptions = []

    # Process each segment
    for i in range(num_segments):
        start_time = i * segment_length
        end_time = min((i + 1) * segment_length, total_length)

        start_sample = int(start_time * expected_sample_rate)
        end_sample = int(end_time * expected_sample_rate)

        segment_waveform = waveform[:, start_sample:end_sample]

        # Process the audio to get the input features
        inputs = processor(segment_waveform.squeeze(0), sampling_rate=expected_sample_rate, return_tensors="pt")

        # Generate the transcription
        with torch.no_grad():
            generated_ids = model.generate(inputs["input_features"], max_length=1000)  # Adjust max_length if needed

        # Decode the transcription
        transcription = processor.batch_decode(generated_ids, skip_special_tokens=True)
        transcriptions.append(transcription[0])
        print(f"Segment {i + 1}/{num_segments} transcribed")

        # Clear memory
        del inputs, segment_waveform, generated_ids
        torch.cuda.empty_cache()

    # Combine all segments
    full_transcription = " ".join(transcriptions)
    print("Full transcription completed")
    if "Titulky vytvořil JohnyX." in full_transcription:
        full_transcription = full_transcription.replace(" Titulky vytvořil JohnyX.", '', 1)
    return full_transcription


def transcribe_english_audio(audio_file_path,segment_length=30):
    """
    nonfinetuned base model for transcribing english audio
    """

    # Load the processor and the model (ensure these names are correct)
    processor = AutoProcessor.from_pretrained("openai/whisper-large-v3")
    model = AutoModelForSpeechSeq2Seq.from_pretrained("openai/whisper-large-v3")
    # Load and preprocess the audio file
    try:
        waveform, sample_rate = torchaudio.load(audio_file_path)
    except FileNotFoundError:
        print(f"Error: Audio file '{audio_file_path}' not found.")
        return None

    # Convert to mono if necessary
    if waveform.shape[0] > 1:
        waveform = torch.mean(waveform, dim=0, keepdim=True)

    # Resample if sample rate doesn't match model's expectation
    expected_sample_rate = processor.feature_extractor.sampling_rate
    if sample_rate != expected_sample_rate:
        waveform = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=expected_sample_rate)(waveform)

    # Calculate the number of segments
    total_length = waveform.shape[1] / expected_sample_rate
    num_segments = int(total_length // segment_length) + 1

    transcriptions = []

    # Process each segment
    for i in range(num_segments):
        start_time = i * segment_length
        end_time = min((i + 1) * segment_length, total_length)

        start_sample = int(start_time * expected_sample_rate)
        end_sample = int(end_time * expected_sample_rate)

        segment_waveform = waveform[:, start_sample:end_sample]

        # Process the audio to get the input features
        inputs = processor(segment_waveform.squeeze(0), sampling_rate=expected_sample_rate, return_tensors="pt")

        # Generate the transcription
        with torch.no_grad():
            generated_ids = model.generate(inputs["input_features"], max_length=1000)  # Adjust max_length if needed

        # Decode the transcription
        transcription = processor.batch_decode(generated_ids, skip_special_tokens=True)
        transcriptions.append(transcription[0])

    # Combine all segments
    full_transcription = " ".join(transcriptions)

    
    
    return full_transcription




# Example usage:
# full_text = transcribe_czech_audio("path/to/audio.wav", "path/to/output.txt")


def create_document(image_path, items_dict, doc_path='output.docx'):
    """
    creation docx proposal document
    """
    # Create a new Document
    doc = Document()

    # Add an image to the document
    doc.add_picture(image_path, width=Inches(5))

    # Add a paragraph for spacing
    doc.add_paragraph()

    # Dictionary with categories and their explanations
    categories_explanations = {
        "Target functionality": items_dict.get("Target functionality", ""),
        "Solution": items_dict.get("Solution", ""),
        "Inputs": items_dict.get("Inputs", ""),
        "Outputs": items_dict.get("Outputs", ""),
        "Requirements and assumptions": items_dict.get("Requirements and assumptions", ""),
        "Human review of outputs": items_dict.get("Human review of outputs", ""),
        "PoC vs production implementation": items_dict.get("PoC vs production implementation", ""),
        "Post-PoC improvements and functionality add-ons": items_dict.get("Post-PoC improvements and functionality add-ons", ""),
        "Infrastructure": items_dict.get("Infrastructure", ""),
        "Limitations": items_dict.get("Limitations", ""),
        "Budget": items_dict.get("Budget", ""),
    }

    # Add a table with two columns
    table = doc.add_table(rows=1, cols=2)

    # Set the widths of the columns
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4)

    # Set the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Explanation'

    # Add the categories and their explanations to the table
    for category, explanation in categories_explanations.items():
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = explanation

    # Save the document
    doc.save(doc_path)
    print(f'Document saved as {doc_path}')

def create_document_2(image_path, items_dict):
    """
    Create a DOCX proposal document and return it as a binary stream.
    """
    # Create a new Document
    doc = Document()

    # Add an image to the document
    doc.add_picture(image_path, width=Inches(5))

    # Add a paragraph for spacing
    doc.add_paragraph()

    # Dictionary with categories and their explanations
    categories_explanations = {
        "Target functionality": items_dict.get("Target functionality", ""),
        "Solution": items_dict.get("Solution", ""),
        "Inputs": items_dict.get("Inputs", ""),
        "Outputs": items_dict.get("Outputs", ""),
        "Requirements and assumptions": items_dict.get("Requirements and assumptions", ""),
        "Human review of outputs": items_dict.get("Human review of outputs", ""),
        "PoC vs production implementation": items_dict.get("PoC vs production implementation", ""),
        "Post-PoC improvements and functionality add-ons": items_dict.get("Post-PoC improvements and functionality add-ons", ""),
        "Infrastructure": items_dict.get("Infrastructure", ""),
        "Limitations": items_dict.get("Limitations", ""),
        "Budget": items_dict.get("Budget", ""),
    }

    # Add a table with two columns
    table = doc.add_table(rows=1, cols=2)

    # Set the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Explanation'

    # Add the categories and their explanations to the table
    for category, explanation in categories_explanations.items():
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = explanation

    # Save the document to a binary stream
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return docx_stream

def create_excel_with_values(notes, file_name='price_estimate.xlsx'):
    """
    creation of excel for the budget
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Price Estimate"

    # Merge cells for title
    ws.merge_cells('A1:D1')
    title_cell = ws['A1']
    title_cell.value = "Price Estimate"
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = Alignment(horizontal="center")
    title_cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    # Add headers with styles
    headers = ["Mandays Min", "Mandays Min + Extra", "Mandays Max", "Mandays Max + Extra"]
    header_fill_min = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")
    header_fill_max = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
        if "Max" in header:
            cell.fill = header_fill_max
        else:
            cell.fill = header_fill_min

    values = [0, 1, 2, 3]

    for col_num, value in enumerate(values, 1):
        value_cell = ws.cell(row=3, column=col_num)

        if value == 0:
            aspect = "Mandays Min"
        elif value == 1:
            aspect = "Mandays Min + Extra"
        elif value == 2:
            aspect = "Mandays Max"
        elif value == 3:
            aspect = "Mandays Max + Extra"

        value_cell.value = mandays_chatgpt(notes,aspect)
        value_cell.alignment = Alignment(horizontal="center")

        if col_num == 2 or col_num == 4:
            value_cell.fill = header_fill_max
        else:
            value_cell.fill = header_fill_min

    # Switch the background colors of B3 and C3
    ws['B3'].fill = header_fill_min
    ws['C3'].fill = header_fill_max

    # Set column widths to fit longer labels
    for col_num, header in enumerate(headers, 1):
        max_length = len(header)
        for row in ws.iter_rows(min_row=2, max_row=3, min_col=col_num, max_col=col_num):
            for cell in row:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
        adjusted_width = (max_length + 2)
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_num)].width = adjusted_width

    # Save the workbook
    wb.save(file_name)
    print(f"Excel file '{file_name}' created successfully.")

def update_excel_file(file_name='price_estimate.xlsx'):
    """
    update excel file with new values
    """
    wb = openpyxl.load_workbook(file_name)
    ws = wb.active

    print("Here are the current values:")
    values = []
    headers = ["Mandays Min", "Mandays Min + Extra", "Mandays Max", "Mandays Max + Extra"]
    for col_num, header in enumerate(headers, 1):
        value = ws.cell(row=3, column=col_num).value
        values.append(value)
        print(f"{header}: {value}")

    # Ask user for each value individually
    for i, (header, value) in enumerate(zip(headers, values)):
        update = input(f"Are you happy with the value for {header} ({value})? (yes/no): ").strip().lower()
        if update == 'no':
            new_value = int(input(f"Enter new value for {header}: "))
            ws.cell(row=3, column=i + 1, value=new_value).alignment = Alignment(horizontal="center")

    wb.save(file_name)
    print(f"Excel file '{file_name}' updated successfully.")



def create_excel_with_values_2(notes):
    """
    creation of excel for the budget
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Price Estimate"

    # Merge cells for title
    ws.merge_cells('A1:D1')
    title_cell = ws['A1']
    title_cell.value = "Price Estimate"
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = Alignment(horizontal="center")
    title_cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    # Add headers with styles
    headers = ["Mandays Min", "Mandays Min + Extra", "Mandays Max", "Mandays Max + Extra"]
    header_fill_min = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")
    header_fill_max = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
        if "Max" in header:
            cell.fill = header_fill_max
        else:
            cell.fill = header_fill_min

    values = [0, 1, 2, 3]

    for col_num, value in enumerate(values, 1):
        value_cell = ws.cell(row=3, column=col_num)

        if value == 0:
            aspect = "Mandays Min"
        elif value == 1:
            aspect = "Mandays Min + Extra"
        elif value == 2:
            aspect = "Mandays Max"
        elif value == 3:
            aspect = "Mandays Max + Extra"

        value_cell.value = mandays_chatgpt(notes, aspect)
        value_cell.alignment = Alignment(horizontal="center")

        if col_num == 2 or col_num == 4:
            value_cell.fill = header_fill_max
        else:
            value_cell.fill = header_fill_min

    # Switch the background colors of B3 and C3
    ws['B3'].fill = header_fill_min
    ws['C3'].fill = header_fill_max

    # Set column widths to fit longer labels
    for col_num, header in enumerate(headers, 1):
        max_length = len(header)
        for row in ws.iter_rows(min_row=2, max_row=3, min_col=col_num, max_col=col_num):
            for cell in row:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
        adjusted_width = (max_length + 2)
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_num)].width = adjusted_width

    # Save the workbook to a BytesIO buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def update_excel_file_2(file_name='price_estimate.xlsx'):
    """
    update excel file with new values
    """
    wb = openpyxl.load_workbook(file_name)
    ws = wb.active

    print("Here are the current values:")
    values = []
    headers = ["Mandays Min", "Mandays Min + Extra", "Mandays Max", "Mandays Max + Extra"]
    for col_num, header in enumerate(headers, 1):
        value = ws.cell(row=3, column=col_num).value
        values.append(value)
        print(f"{header}: {value}")

    # Ask user for each value individually
    for i, (header, value) in enumerate(zip(headers, values)):
        update = input(f"Are you happy with the value for {header} ({value})? (yes/no): ").strip().lower()
        if update == 'no':
            new_value = int(input(f"Enter new value for {header}: "))
            ws.cell(row=3, column=i + 1, value=new_value).alignment = Alignment(horizontal="center")

    # Save the workbook to a BytesIO buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer





decoder_prompt_1 = """---

**Task:**

You are given raw OCR outputs and your task is to manually correct and improve them based on contextual understanding, pattern recognition, and common phrases. Follow the steps outlined below to transform the OCR text into coherent and accurate corrected text. Use the provided examples as guides.

**Steps:**

1. **Review the OCR Output**: Carefully read through the OCR text. Note any obvious errors, misinterpretations, and nonsensical phrases.
2. **Identify Common Phrases**: Use your knowledge of the subject matter to identify and replace common phrases and terms that are likely misinterpreted by the OCR process.
3. **Contextual Understanding**: Apply context to make educated guesses about unclear sections. Use related terminology if you recognize the topic.
4. **Manual Correction**: Correct the OCR text line by line, ensuring consistency and readability.
5. **Provide the Final Corrected Text**: Rewrite the entire text with your corrections, ensuring it is coherent and accurate.

**Example 1:**

**OCR Output:**
```
‘ el | 3.9 ‘ ee ae e ‘
mf | : | 4g mapa gig
ts |
| / ; Y a .
a edad ws Bak - See OF i ® 2
ee Ns Lule adit A ea DPB *~,
| _- xe low chty KL v oil pasted NY 4 ~
wae ye —pvyy GUCOCLOUV | _y,—
os ay es rere Fe aE  Dayfats 2 Hs poe Witter 4 oy
| - Cee Low S 2 li lele” to ugpiinila A a — a
_e chifltrr rl Tip? Vy ee
= sevlete ca bl 7 et L i Money = PV Oe —_—
“ae a oS! ne "
a ae
Te
eae ee ee A ™
a ee eee
Ret A a 2 a
: 2, he Se § gu-Pve see ~»
GH) C loru lb —C eve” 7 CPE J? a?" a)
" Hotty yorptatts) pivalela C bdai =e a )
2 Jxh probthe fv, eG 4 ee C7 po the lead) : ”
t piy'pe, her) He praftpely Nhl le ~ YPH/ Cb eor 5 ”
+ leet “fap pile Geety Pe 20. — ZELW £ Brel,” Frais ©
= AMMEN 9 vetin Ce Fb Kk OX be Pr” Ponty lece B46,
—  - aml) tea 9 Eee) Oy fv ‘
-/ PL parley p AB Fda EL acto 2
maa gy lee Ca, Yee Ulipt! briny brad ~ ee Clee yy ”
OO tel att, Sow ter told
ET 2. - 3
| » Mavia Ak, PoclGre - 4th ftta fy fe hbey we ?
el OO ORS OP RO CR OU Cr
| » Lat Plt na ttle 4 ees s ihe) ee ., ?
| ee - |
Se ey ee 1. ?
_——— : bGly" FS VPbya (a SS. 35 aa
Be 8) hes = eed
Be 1) che tn [bitin [prin Jug, 7 Pe
| Hla ab chs Ville cs = |
q So ee Ma veh mR ae $A | |
— =the 4 Af. Seo A a gz
```

**Common Phrases Identified:**
- "Cells are the basic unit of life."
- "Organelles within cells carry out specific functions."
- "Membrane dynamics and protein transport are crucial."

**Contextual Understanding:**
This text appears to be a set of notes on cell biology. Key concepts include cellular functions, membrane transport, energy production, protein synthesis, and cell division.

**Manual Correction:**
```
Cell Biology Notes:

- Cells are the basic unit of life.
- Organelles within cells carry out specific functions.
- Membrane dynamics and protein transport are crucial.
- Study how cells communicate and interact with their environment.
- Understand the impact of genetic mutations on cellular function.

Mechanisms of Cellular Functions:

1. Membrane Transport:
   - Diffusion
   - Osmosis
   - Active Transport

2. Energy Production:
   - Cellular Respiration
   - ATP synthesis

3. Protein Synthesis:
   - Transcription
   - Translation

4. Cell Division:
   - Mitosis
   - Meiosis

Research Topics:
- Impact of environmental factors on cell behavior.
- Role of stem cells in tissue regeneration.
- Advances in cancer treatment targeting cellular mechanisms.
- Development of artificial cells and their potential applications.

Laboratory Techniques:
- Microscopy for cell imaging.
- Gel electrophoresis for DNA analysis.
- PCR for amplifying genetic material.
- Western blot for protein detection.

Challenges in Cell Biology:
- Understanding complex signaling pathways.
- Overcoming drug resistance in cancer cells.
- Ethical considerations in genetic manipulation.
- Integrating multidisciplinary approaches for comprehensive research.

Future Directions:
- Personalized medicine based on individual cellular profiles.
- Engineering cells for therapeutic purposes.
- Exploring the potential of nanotechnology in cell biology.
- Enhancing tissue engineering and regenerative medicine.

This document provides a comprehensive overview of essential concepts in cell biology, focusing on cellular functions, research areas, and laboratory techniques. It aims to support students and researchers in advancing their understanding and contributing to the field.
```

**Example 2:**

**OCR Output:**
```
phy | cloth << teliig OF Nompyeble, > HF auf point fice
```

**Common Phrases Identified:**
- "physical effects"
- "structure of nanoparticles"
- "different geometries"

**Contextual Understanding:**
This text discusses the physical properties and structures of nanoparticles.

**Manual Correction:**
```
physical effects & structure of nanoparticles -> different geometries
```

**Example 3:**

**OCR Output:**
```
re, nah le i lo Shells AF
bot, dahanay VeYlalan A procbusthin, be ty 5 /o role
```

**Common Phrases Identified:**
- "KPIs"
- "vzhledem k produktivitě"
- "kolem 20"

**Contextual Understanding:**
This text is related to key performance indicators (KPIs) for restaurants in terms of productivity.

**Manual Correction:**
```
KPIs restaurací vzhledem k produktivitě by bylo kolem 20
```

**Example 4:**

**OCR Output:**
```
1 — chin fo Fad Ghilu& 2° pou pete, Ee Yateupbhp, Bre -Hoylee,
```

**Common Phrases Identified:**
- "similar to hard structures of the body"
- "Biocompatible"
- "Bio-degradable"

**Contextual Understanding:**
This text seems to be discussing materials that are similar to hard structures of the body and their biocompatibility and biodegradability.

**Manual Correction:**
```
1 - similar to hard structures of the body, ECA, Biocompatible, Bio-degradable,
```

**Example 5:**

**OCR Output:**
```
ESAB - SRC (Student report)
    - big announcement today! (They survived first round!)
    - new leaders and editors

    LP - friendly meetings schedule
    - overlap in student clubs - faculty meetings, research work.

    Joy:
    RT: What is it made of? How do they get to the target?
    1. what is it made of?
    2. how can they "hit" the target?
        - oxygen
        - heat
        - cells
        - organelles

    3. how influence the shape?
    4. how obtain the dry molecules? Where to - including labs culture?

    5. How is bigger affecting?

1. similar to hard structures of the body, ECA, biocompatible, bio-degradable.
    - flexible -> other materials, semi-solid, ...
    - IP -> includes NDA, adherence crystals

2. containment + retrieval - can it be degraded & kept stable

3. therapeutic cargo
    physical effects & working of nanoparticles -> different geometries
    of the structures. Untested NDs, more hyperactive states.
    can carry universal cargoes. They can influence properties - charge, ...
    - stability, when building ix.
    - affinity to bound -> cells
    - networks connecting particles
    - solubility? tagging; it can be added to other containment layers
    - combination -> things are possible
```

**Common Phrases Identified:**
- "summary memo"
- "Store costs"
- "big announcement today"
- "leaders and editors"
- "friendly meetings schedule"

**Contextual Understanding:**
This text appears to be a summary memo discussing store costs, announcements, leadership changes, and meeting schedules.

**Manual Correction:**
```
NS - summary memo:
UE Store costs 527

- big announcement today! (They survived the first round!)
- new leaders and editors

LP - friendly meetings schedule
- overlap in student clubs - faculty meetings, research work.

Joy:
RT: What is it made of? How do they get to the target?
1. what is it made of?
2. how can they "hit" the target?
    - oxygen
    - heat
    - cells
    - organelles

3. how influence the shape?
4. how obtain the dry molecules? Where to - including labs culture

?

5. How is bigger affecting?

1. similar to hard structures of the body, ECA, biocompatible, bio-degradable.
    - flexible -> other materials, semi-solid, ...
    - IP -> includes NDA, adherence crystals

2. containment + retrieval - can it be degraded & kept stable

3. therapeutic cargo
    physical effects & working of nanoparticles -> different geometries
    of the structures. Untested NDs, more hyperactive states.
    can carry universal cargoes. They can influence properties - charge, ...
    - stability, when building ix.
    - affinity to bound -> cells
    - networks connecting particles
    - solubility? tagging; it can be added to other containment layers
    - combination -> things are possible
```

**Example 6:**

**OCR Output:**
```
Dans /
it) «Hen hor) Love)” -
‘ tl Yin WAM) 7 Uply Kyl’) Pewee >
“a 3 pm 44 Lr Aww L
0 a Ltt ae poe 4a Mery” Veo y Af
taht 1 HR bu he
| 1 Melee,” gkyqpuy > po pve cede — ued bet;
| — eye fy HEL ee 4 O° 10 2?" / CO
ep ep Cb hela 64 weg) fan V Jt aetiy)) Pvt Xhapyg:
- “s VY ne ve4 Bp a Dlal BL Z ele3
a | eee Whe! Aq & cbf
“a |  Sewabuy veetwou” $ a. 7? loca! A> ces
<— | ae
terme 2A M7 teov20ty ren er?
oa = | > chip X ie 49 &> vane” 2A Le ett 2e neds,
_ 41) col erer (-tetilin’ — —~)eveliy peo Je Abe
_ a | Ct & Desf afer) — periend [0h ee Area”
eo GR A ya
. a e , — phe wotivelitits pot sd wlerdy
= ra | © eo = Tet a) Ae pois at wh ply stip yz) bee Bebe
| Spey eee — 2 AHOLL 2 Ct eh belt) cb,
5 - , 5 a — +> J
— #) lay lenis) oles Jweb SPALL CEM Vo).
_ —_— | : ] peLL: ee ey,
a EEE
_ 3
7 wveth ot Pay Fvtery os , Tu REV LA by ee 7
: | — Liha, Corp Av nich That — = ~W/s tee ete ley py
— — yutlrybmn OC CPAcu ct Aye pia pols / Neeerbtae | 2 Peapiyon ce
—> +) tealev pubelis’ Amit s —Aviey 4 pole gple =
a a = a #62 fa dponiia Me Plain "bn Hoe ge oe
. — cof i enters # 464) 0 bor Gye tt ) Bese
I
a een PP eer ene
“ | 5 ae
= | cuenta. => Wt!) pute eas ee Ly, a
A | nner’ = 40 2= a2”
| eee Aa ee
7, a Tone NT th, Zt Ely fe ,, okt akce;7t tentonst
— 2 tse ag | fp Mpa" beeen © ak) Cevtge
on SF | aah vhaled pre roe) yp GODT ee pes
| gecmig) EU Wir ¢ wh file pei! Per” ee nw *
ee aa bhsse Pewsey 240) = ee
. ee — j = bbin “yh WY, per) pa Puce OH
, i 2. eeosee Uier ash pv ee en eh We
IG .! 2 Gila 4 folk foe jt ee, a oe dn
| le Saat ee Se
oi — ee sete) ‘pele 4? Loy typ la foe see? 00 Efe ) mahiaay ve Sreplfley th, 5
Ce FE OT oi 6mm Aes
pay PFT) G8
```

**Common Phrases Identified:**
- "Understanding Love Languages"
- "Importance of Active Listening"
- "Different Communication Styles"

**Contextual Understanding:**
This text appears to be notes on communication styles and love languages, focusing on relationship dynamics and effective communication techniques.

**Manual Correction:**
```
Random Notes:

- Discussion on various love languages and their impact on relationships.
- Highlight the importance of active listening and empathy in communication.
- Explore different communication styles and their effectiveness.

Key Points:

1. Understanding Love Languages:
   - Words of Affirmation
   - Quality Time
   - Receiving Gifts
   - Acts of Service
   - Physical Touch

2. Importance of Active Listening:
   - Fully focus on the speaker
   - Avoid interrupting
   - Show interest through body language

3. Different Communication Styles:
   - Assertive
   - Aggressive
   - Passive
   - Passive-Aggressive

Research Topics:
- Impact of love languages on relationship satisfaction.
- Role of empathy in conflict resolution.
- Effectiveness of different communication styles in various situations.

Future Directions:
- Developing tools for better communication in relationships.
- Exploring cultural differences in love languages and communication styles.
- Creating educational programs to enhance interpersonal skills.
```

---

By following these steps and examples, you can accurately correct OCR outputs and transform them into coherent, readable text. just return the readable text """

decoder_prompt_2 = """You are tasked with converting OCR outputs to the intended texts. Here are the observed patterns and examples to guide you:
Observed Patterns:

    Correction of Typographical Errors:
        Replace incorrect or misinterpreted characters with the correct ones (e.g., "ECO" to "ECA").
        Correct spelling errors and ensure proper terminology is used.

    Rephrasing for Clarity:
        Ensure sentences are clear and meaningful, adjusting phrases as necessary (e.g., "cy endosomal stage" to "endosomal stage").

    Maintaining Context and Meaning:
        Preserve the original context and meaning of the text while making corrections.

    Formatting and Structure:
        Ensure the text is well-structured, with proper headings, subheadings, and bullet points as seen in examples.

Examples:

    Example 1:
        OCR Output: IDB. Koronen ? nadom, ?odnala ?jackle ?, Ko ?o salon 20%
        Intended Output: KPIs restaurací vzhledem k produktivitě by bylo kolem 20

    Example 2:
        OCR Output: 1- similar to hard structures of the body, ECM, biodegradable, bio-compatible
        Intended Output: 1 - similar to hard structures of the body, ECA, Biocompatible, Bio-degradable

    Example 3:
        OCR Output: ESOR - SRR (Endoven sequels)

            be endosome till? (cy endosomal stage) Possible!
            -> endosomal
            -> interactions & elders
            -> IP - towards the family - Soldiers,
            -> Climate no endosomal cages - towards the body, + tasks in protein work.

            Overtake no endosomal cages - towards the body, + tasks in protein work.

            what is it made of? How do they go to reading?

            how can they “visit the target?
                oxygen
                lister
                cell
                organelle

            how influence shipping?

            how clear the dry molecules? Where to – Including release

            cells,

            How bioengin donicity?

    1 - similar to head clusters of the body, ECO, (biocomaptible), bio-clustering IP -subdivide NAD -> nanomaterials, sew-intellect
    printer -
    glucoses -> crytals

    2 - containment + motivating - can it degrade (8 temp./adult stable

    4 - Therapeutic cargo
    physical signifying & meaning of nanomaterials -> differential quantities
    of the endosome - Untitled MCPs, room temperature states.
    Can carry universal cargoes. They can influence -> charge/ peroxide -> when burning X.

        Intended Output:
        ESOR - SRR (Endoven sequels)

            The endosome tail? (cy endosomal stage) Possible!
            -> endosomal
            -> interactions & leaders
            -> IP - towards the family - Soldiers,
            -> Climate on endosomal cages - towards the body, + tasks in protein work.

                Overtake on endosomal cages - towards the body, + tasks in protein work.

                What is it made of? How do they go to reading?

                How can they “visit the target?
                    oxygen
                    lister
                    cell
                    organelle

                How influence shipping?

                How clear the dry molecules? Where to – Including release

                cells,

                How bioengine donicity?

        1 - Similar to hard clusters of the body, ECA, (biocompatible), bio-clustering IP - subdivide NAD -> nanomaterials, sew-intellect
        printer -
        glucoses -> crystals

        2 - containment + motivating - can it degrade (8 temp./adult stable)

        4 - Therapeutic cargo
        Physical effects & working of nanoparticles -> different geometries
        of the endosome - Untested NDs, more hyperactive states.
        Can carry universal cargoes. They can influence -> charge/ peroxide -> when burning X.

    Example 4:

        OCR Output: ESAB - SRC (Student report)

        big announcement today! (They survived the first round!)
        new leaders and editors

  LP - friendly meetings schedule

overlap in student clubs - faculty meetings, research work.

vbnet

          Joy:
          RT: What is it made of? How do they get to the target?

        what is it made of?
        how can they "hit" the target?
        oxygen
        heat
        cells
        organelles

        how influence the shape?
        how obtain the dry molecules? Where to - including labs culture?

        How is bigger affecting?

        similar to hard structures of the body, ECA, biocompatible, bio-degradable.
        flexible -> other materials, semi-solid, ...
        IP -> includes NDA, adherence crystals

        containment + retrieval - can it be degraded & kept stable

        therapeutic cargo
        physical effects & working of nanoparticles -> different geometries
        of the structures. Untested NDs, more hyperactive states.
        can carry universal cargoes. They can influence properties - charge, ...
        stability, when building ix.
        affinity to bound -> cells
        networks connecting particles
        solubility? tagging; it can be added to other containment layers
        combination -> things are possible

        Intended Output:
        ESAB - SRC (Student report)

        Big announcement today! (They survived the first round!)
        New leaders and editors

        LP - friendly meetings schedule

        Overlap in student clubs - faculty meetings, research work.

        Joy:
        RT: What is it made of? How do they get to the target?

        What is it made of?
        How can they "hit" the target?
            oxygen
            heat
            cells
            organelles

        How influence the shape?
        How obtain the dry molecules? Where to - including labs culture?

        How is bigger affecting?

        Similar to hard structures of the body, ECA, biocompatible, biodegradable.
        Flexible -> other materials, semi-solid, ...
        IP -> includes NDA, adherence crystals

        Containment + retrieval - can it be degraded & kept stable?

        Therapeutic cargo
        Physical effects & working of nanoparticles -> different geometries
        of the structures. Untested NDs, more hyperactive states.
        Can carry universal cargoes. They can influence properties - charge, ...
        Stability, when building ix.
        Affinity to bound -> cells
        Networks connecting particles
        Solubility? Tagging; it can be added to other containment layers
        Combination -> things are possible

Use these patterns and examples to accurately convert OCR outputs to the intended texts. Ensure clarity, correct terminology, and proper formatting in the final text."""


decoder_prompt_3 = """
Observations:

    Character Misrecognition: OCR often misreads characters, e.g., "hand" instead of "hard," "ECB" instead of "ECA."
    Word Misrecognition: OCR sometimes misinterprets words entirely, e.g., "preorbital swelling" instead of a completely different phrase.
    Formatting Issues: Punctuation and capitalization may be inconsistent or missing.
    Partial Phrases: OCR may capture incomplete or jumbled phrases.

Decoder Prompt:

Decoder Instructions for Converting OCR Outputs to Intended Texts:

    Character Correction:
        Review common OCR misrecognitions and correct them. For example, replace "hand" with "hard," "ECB" with "ECA," and ensure biocompatible is capitalized correctly.

    Word Correction:
        Look for context clues to determine correct words. For instance, if the OCR output mentions "preorbital swelling," check the intended text for a related but contextually accurate phrase.
        Correct specific terms frequently misinterpreted by OCR, such as "biotics" to "KPIs."

    Consistency in Formatting:
        Ensure proper capitalization of specific terms, especially at the beginning of sentences and for proper nouns.
        Correct punctuation as needed, adding commas, periods, and hyphens where necessary.

    Contextual Understanding:
        Use the context provided by the surrounding text to make sense of ambiguous or unclear words. For example, if "IP" is followed by "subjects NDA glucose," the intended text may refer to "IP includes NDA, adherence crystals."

    Structural Corrections:
        Ensure that bullet points, numbering, and section headings match the intended structure. For example, if a list is present, ensure each item is correctly numbered and formatted.

Example Conversions:

OCR Output:
"1- similar to hand structures of the body, ECB, biocompatible, bio-degradable"

Intended Text:
"1 - similar to hard structures of the body, ECA, Biocompatible, Bio-degradable,"

OCR Output:
"how can they "hurt the target?"

    oxygen
    heat
    cell
    organelles"

Intended Text:
"how can they "hit" the target?

    oxygen
    heat
    cells
    organelles"

Decoder Prompt:

vbnet

Given the OCR output, apply the following corrections to produce the intended text:

1. Correct common character misrecognitions, such as replacing "hand" with "hard" and "ECB" with "ECA."
2. Use context to determine correct words, e.g., replacing "biotics" with "KPIs."
3. Ensure proper capitalization and punctuation throughout the text.
4. Maintain the intended structure, including bullet points, numbering, and section headings.

Example:
OCR Output: "1- similar to hand structures of the body, ECB, biocompatible, bio-degradable"
Intended Text: "1 - similar to hard structures of the body, ECA, Biocompatible, Bio-degradable,"

By following this structured prompt, the task of converting OCR outputs to the intended texts can be approached methodically, improving accuracy and efficiency.
"""
decoder_prompt_4 ="""
OCR Outputs:

    "Bt, ou hanay” Veblen A prodbud bj te ax, regray” VEU lila, A proba be 4 hf Yolen ae"
    "7 — chine fo Ped chil’ 0 Je hee, BCb2 ) Plotepdhp, Bre -Chydee,"
    "Pie AL Vrs, wy — Zhe see I To ee Cae a ile : : “ a as AY / ‘ BS IM a Lege AS ae SOP ime oe Bac” cles gah Bt Pb (004. "1 Opp, 0 “oboe OR - jirswits pb OG a — EN rn = 4) Whe 1 Yh YR re (ome pa wy @) Foe can Yong “UE PA PGT So a se eS “2 es ee Le ee a _ Ow hin ett = aa 4) be A bere Ye ts waleuteg F Where eo —n/o mt LO Meta, © = — SS hee w hovtgn ve, Cle Wepes © oo e 4 — ety ln J obilut 06 Ju per. , C07 plotdarhhh, Bre > iad - hb , Of}ey r9
    Yl £000"-$2)27 ; = : 5 7 —s nnehbe, vith bh ony h Friucg o rhe . 7 h- O11 a ah - Ye Ges bxbe A hip cll sible = " , Lyiy| clube & e1 ben et hepwpbble, > AS Gu Zttintti ii = * ; bt Sou ln , od NPs, Veve, Jayotns lable, ast lin Can, biases, Cer Thy Se aes le ale =Cb ty aad i a fF 1, of Stich a) aM ee aa . a — nepal counulia, chs be a 2 - 2 bfaliuih , Vay Fy s Vt tty $¢ Gili A be pby Cth [Cpt rene / lye, ;"

Decoder Prompt:

Decoder Instructions for Converting OCR Outputs to Intended Text:

    Identify Common Patterns:
        Review all provided OCR outputs for common patterns and discrepancies.
        Note that OCR outputs may vary significantly in character recognition and word formation.

    Correct Common Misrecognitions:
        Replace commonly misrecognized characters with the correct ones, e.g., replace "hanay" with "hand," "BCb2" with "ECA."
        Fix punctuation and capitalization inconsistencies.

    Use Context Clues:
        Use the surrounding context to determine the intended words and phrases. For example, if "Bt, ou hanay" is repeated in another OCR output as "chin fo Ped chil'," consider the possibility of the intended text being related to "similar to hard structures."

    Reconstruct Fragmented Texts:
        Merge fragmented phrases and sentences based on the logical flow of the text. For instance, "0 Je hee, BCb2 ) Plotepdhp, Bre -Chydee," can be reconstructed with the intended scientific terminology and structure.

    Ensure Consistency and Coherence:
        Ensure the final text is coherent and consistent in structure, grammar, and meaning.

Example Conversion:

OCR Output 1:
"Bt, ou hanay” Veblen A prodbud bj te ax, regray” VEU lila, A proba be 4 hf Yolen ae"

OCR Output 2:
"7 — chine fo Ped chil’ 0 Je hee, BCb2 ) Plotepdhp, Bre -Chydee,"

OCR Output 3:
"Pie AL Vrs, wy — Zhe see I To ee Cae a ile : : “ a as AY / ‘ BS IM a Lege AS ae SOP ime oe Bac” cles gah Bt Pb (004. "1 Opp, 0 “oboe OR - jirswits pb OG a — EN rn = 4) Whe 1 Yh YR re (ome pa wy @) Foe can Yong “UE PA PGT So a se eS “2 es ee Le ee a _ Ow hin ett = aa 4) be A bere Ye ts waleuteg F Where eo —n/o mt LO Meta, © = — SS hee w hovtgn ve, Cle Wepes © oo e 4 — ety ln J obilut 06 Ju per. , C07 plotdarhhh, Bre > iad - hb , Of}ey r9
Yl £000"-$2)27 ; = : 5 7 —s nnehbe, vith bh ony h Friucg o rhe . 7 h- O11 a ah - Ye Ges bxbe A hip cll sible = " , Lyiy| clube & e1 ben et hepwpbble, > AS Gu Zttintti ii = * ; bt Sou ln , od NPs, Veve, Jayotns lable, ast lin Can, biases, Cer Thy Se aes le ale =Cb ty aad i a fF 1, of Stich a) aM ee aa . a — nepal counulia, chs be a 2 - 2 bfaliuih , Vay Fy s Vt tty $¢ Gili A be pby Cth [Cpt rene / lye, ;"

Intended Text:

vbnet

how can they "hit" the target?
- oxygen
- heat
- cells
- organelles

how influence the shape?
how obtain the dry molecules? Where to - including labs culture?

How is bigger affecting?

similar to hard structures of the body, ECA, biocompatible, bio-degradable.
flexible -> other materials, semi-solid, ...
IP -> includes NDA, adherence crystals

containment + retrieval - can it be degraded & kept stable

therapeutic cargo
physical effects & working of nanoparticles -> different geometries
of the structures. Untested NDs, more hyperactive states.
can carry universal cargoes. They can influence properties - charge, ...
stability, when building ix.
affinity to bound -> cells
networks connecting particles
solubility? tagging; it can be added to other containment layers
combination -> things are possible.

By following these steps, you can accurately decode the OCR outputs to the intended texts.
"""

def extract_text_from_image(api_key, uploaded_file):
    """
    Translate image to base64 and then perform OCR on it.
    """
    # Function to encode the image
    def encode_image(file):
        return base64.b64encode(file.read()).decode('utf-8')

    # Convert PIL Image to BytesIO object if necessary
    if isinstance(uploaded_file, Image.Image):
        buffered = BytesIO()
        uploaded_file.save(buffered, format="JPEG")
        buffered.seek(0)
        base64_image = encode_image(buffered)
    else:
        base64_image = encode_image(uploaded_file)

    # Define headers and payload for the API request
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    payload = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "you are given a handwritten text image which you are tasked to convert to text, just provide the text nothing else no explanation"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 300
    }

    # Make the API request
    response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)

    # Extract and return the text from the response
    if response.status_code == 200:
        response_data = response.json()
        if "choices" in response_data:
            return response_data["choices"][0]["message"]["content"]
        else:
            return "Error: Unexpected response format."
    else:
        return f"Error: {response.status_code} - {response.text}"

def ask_chatgpt(question):
    """
    for infrancing the bussnies proposal
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "you are tasked with helping bussnies proposal just create the desired part and don't write anything else ."},
            {"role": "user", "content": question},
        ]
    )
    answer = response['choices'][0]['message']['content']
    return answer


def mandays_chatgpt(notes, aspect):
    """
    for guessing mandays for the given aspect for the excel
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"you are tasked with guesssing {aspect} for price estimate for AI startup, you will be only recive notes and just return  a integer nothing else"},
            {"role": "user", "content": notes},
        ]
    )
    answer = response['choices'][0]['message']['content']
    return answer


def checker(item, part):
    """
    AI checking it's repsonse is correct or not and recommended changes
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "you are tasked with checking if provided text is suitable for given part of bussnies proposal and if you think changes are should made what would you recommend, dont provide a revised version just provide your assessment and reccomendations keep it as short as poosible "},
            {"role": "user", "content": "part of bussnies proposal:" + item + "provided text:" + part },
        ]
    )
    answer = response['choices'][0]['message']['content']
    return answer

def decoder(ocr_output,decoder_prompt):
    """
    for translating ocr output to text
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o",  # or "gpt-3.5-turbo" if you prefer to use that model
        messages=[
            {"role": "system", "content": "You are a tasked with creating normal text from a ocr output from a very bad handwriting  you are just supposed to only return translated text here instructions on how to translate ocr output to text:" + decoder_prompt},
            {"role": "user", "content": ocr_output},
        ]
    )
    answer = response['choices'][0]['message']['content']
    return answer


def process_items(items_dict, notes):

    """
    AI generates nessercy parts of proposal and assign it to the items_dict user can change them via prompts
    """
    print("AI generated Use_case_description :")
    Use_case_description = ask_chatgpt("generate use case description from given notes: "+notes)
    print(f"Use_case_description: {Use_case_description}")
    response = input(f"would you like to enter use case description if so enter y ?")
    if response == 'y':
      Use_case_description =  input(f"Enter desired use case description")

    for item in items_dict:
        task_successful = False
        desired_changes = ""
        while not task_successful:
            prompt = f"{desired_changes} generate {item} for business proposal for {Use_case_description} from : {notes} "
            value = ask_chatgpt(prompt)
            print(f"Assigned value for {item}: {value}")
            print("******************************************************************************************************************************")
            opion_ai = checker(item, value)
            print(f"As AI my opinion on this part of proposal is : {opion_ai}")

            response = input("Happy with output? (y/n): ").strip().lower()
            if response == 'y':
                task_successful = True
                items_dict[item] = value
                print(f"Value for {item} confirmed: {value}\n")
            else:
                desired_changes = input(f"Enter desired changes for {item}: ")

def preprocess_handwritten_image(uploaded_file):
    """
    Preprocess the uploaded handwritten image file for OCR and return the preprocessed image.
    """
    # Load the image from the uploaded file
    image = Image.open(uploaded_file).convert('L')
    image_np = np.array(image)

    # Check if the image is loaded properly
    if image_np is None:
        st.error("Error: Unable to load image.")
        return None

    # Step 1: Noise reduction
    denoised_image = cv2.medianBlur(image_np, 1)

    # Step 2: Binarization
    _, binary_image = cv2.threshold(denoised_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Convert the binary image back to PIL Image format for OCR
    preprocessed_image = Image.fromarray(binary_image)

    return preprocessed_image


def main():
    st.title("Proposal Droid_01")

    openai_api_key = st.text_input("Enter your z API key:", type="password", key="api_key_input")
    if openai_api_key:
        # Initialize OpenAI API with the provided key
        openai.api_key = openai_api_key

    if "notes" not in st.session_state:
        st.session_state.notes = None

    if "items_dict" not in st.session_state:
        st.session_state.items_dict = {
            "Target functionality": None,
            "Solution": None,
            "Inputs": None,
            "Outputs": None,
            "Requirements and assumptions": None,
            "PoC vs production implementation": None,
            "Human review of outputs": None,
            "Post-PoC improvements and functionality add-ons": None,
            "Infrastructure": None,
            "Limitations": None
        }

    if "current_item_index" not in st.session_state:
        st.session_state.current_item_index = 0

    if "notes_finalized" not in st.session_state:
        st.session_state.notes_finalized = False

    if "use_case_finalized" not in st.session_state:
        st.session_state.use_case_finalized = False

    if "use_case_description" not in st.session_state:
        st.session_state.use_case_description = ""

    file_type = st.selectbox("Select the type of notes you want to upload:", ["Image", "Voice Recording"])
    uploaded_file = st.file_uploader("Upload your file", type=["png", "jpg", "jpeg", "wav", "mp3", "m4a"])

    if uploaded_file is not None:
        if file_type == "Image":
            st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)
            preprocess = st.checkbox("Preprocess Image")
            preprocessed_image = uploaded_file

            if preprocess:
                preprocessed_image = preprocess_handwritten_image(uploaded_file)
                if preprocessed_image is not None:
                    st.image(preprocessed_image, caption="Preprocessed Image", use_column_width=True)

            ocr_method = st.selectbox("Select OCR method:", ["Method 1", "Method 2", "Method 3"])
            if st.button("Extract Text from Image"):
                if preprocessed_image is not None:
                    image_to_process = preprocessed_image
                else:
                    image_to_process = uploaded_file

                if ocr_method == "Method 1":
                    st.session_state.notes = extract_text_from_image(openai_api_key, image_to_process)
                elif ocr_method == "Method 2":
                    #ocr_output = tes_ext(image_to_process)
                    st.session_state.notes = extract_text_from_image(openai_api_key, image_to_process)
                elif ocr_method == "Method 3":
                    st.session_state.notes = extract_text_from_image(openai_api_key, image_to_process)
                    pass

            if 'notes' in st.session_state:
                use_decoder = st.checkbox("Use decoder prompt?")
                if use_decoder:
                    decoder_prompt = st.selectbox("Select decoder prompt:", ["Prompt 1", "Prompt 2", "Prompt 3", "Prompt 4"])
                    if st.button("Apply Decoder Prompt"):
                        if decoder_prompt == "Prompt 1":
                            decoder_prompt = decoder_prompt_1
                        elif decoder_prompt == "Prompt 2":
                            decoder_prompt = decoder_prompt_2
                        elif decoder_prompt == "Prompt 3":
                            decoder_prompt = decoder_prompt_3
                        elif decoder_prompt == "Prompt 4":
                            decoder_prompt = decoder_prompt_4
                        st.session_state.notes = decoder(st.session_state.notes, decoder_prompt)
                        st.write(st.session_state.notes)
                else:
                    st.write(st.session_state.notes)
                    # Uncomment and define ocr_method_3 if available
                    # elif ocr_method == "Method 3":
                    #     st.session_state.notes = ocr_method_3(image_to_process)

        elif file_type == "Voice Recording":
            st.audio(uploaded_file, format="audio/wav" if uploaded_file.name.endswith('.wav') else "audio/mp3")
            language = st.radio("Select the language of the audio:", ("Czech", "English"))

            if st.button("Transcribe Audio"):
                if language == "Czech":
                    with st.spinner('Transcribing audio...'):
                        st.session_state.notes = transcribe_czech_audio(uploaded_file)
                elif language == "English":
                    with st.spinner('Transcribing audio...'):
                        st.session_state.notes = transcribe_english_audio(uploaded_file)

                st.success("Transcription completed!")

    if st.session_state.notes:
        st.write("Transcribed Notes:")
        st.write(st.session_state.notes)

        if st.button("Edit Notes"):
            st.session_state.edit_mode = True
        if "edit_mode" in st.session_state and st.session_state.edit_mode:
            edited_notes = st.text_area("Edit your notes here:", st.session_state.notes, key="edit_notes_input")
            if st.button("Save Edited Notes"):
                st.session_state.notes = edited_notes
                st.session_state.edit_mode = False
                st.success("Notes updated!")
        if st.button("Use Notes As Is"):
            st.session_state.notes_finalized = True
            st.success("Using transcribed notes as is.")

    if st.session_state.notes_finalized and not st.session_state.use_case_description:
        st.session_state.use_case_description = ask_chatgpt("generate very short use case description from given notes: " + st.session_state.notes)

    if st.session_state.use_case_description:
        st.write("Use Case Description:")
        st.write(st.session_state.use_case_description)
        st.session_state.use_case_description = st.text_area("Edit Use Case Description:", st.session_state.use_case_description, key="use_case_desc_input")
        if st.button("Finalize Use Case Description"):
            st.session_state.use_case_finalized = True

    if st.session_state.use_case_finalized:
        keys = list(st.session_state.items_dict.keys())

        if st.session_state.current_item_index < len(keys):
            current_key = keys[st.session_state.current_item_index]

            if st.session_state.items_dict[current_key] is None:
                prompt = f" generate a short 2-4 sentence {current_key} for business proposal from : {st.session_state.notes} "
                value = ask_chatgpt(prompt)
                st.session_state.items_dict[current_key] = value
                opion_ai = checker(st.session_state.items_dict[current_key], value)
            else:
                opion_ai = ""

            st.write(f"{current_key}: {st.session_state.items_dict[current_key]}")
            st.write(f" AI Opinion for {current_key}: \n {opion_ai}")

            action = st.selectbox(f"Select action for {current_key}:", ["Accept", "Edit", "Add"], key=f"action_select_{current_key}")

            if action == "Accept" and st.button(f"Confirm {current_key}", key=f"confirm_{current_key}"):
                st.session_state.current_item_index += 1
                st.success(f"{current_key} accepted!")

            elif action == "Edit":
                edited_value = st.text_input(f"Edit {current_key}:", st.session_state.items_dict[current_key], key=f"edit_{current_key}")
                if st.button(f"Save {current_key}", key=f"save_{current_key}"):
                    st.session_state.items_dict[current_key] = edited_value
                    st.session_state.current_item_index += 1
                    st.success(f"{current_key} updated!")
                    opion_ai = checker(st.session_state.items_dict[current_key], edited_value)

            elif action == "Add":
                additional_value = st.text_input(f"Add to {current_key}:", key=f"add_{current_key}")
                if st.button(f"Add to {current_key}", key=f"add_confirm_{current_key}"):
                    prompt_for_reg = f"{additional_value} generate {st.session_state.items_dict[current_key]} for business proposal for {st.session_state.use_case_description} from : {st.session_state.notes} "
                    st.session_state.items_dict[current_key] = ask_chatgpt(prompt_for_reg)
                    st.session_state.current_item_index += 1
                    st.success(f"Added to {current_key}!")

        if st.session_state.current_item_index >= len(keys):
            st.success("All items processed!")
            docx_buffer = create_document_2("Screenshot 2024-05-23 171557.png", st.session_state.items_dict)
            st.download_button(
                label="Download DOCX",
                data=docx_buffer,
                file_name='output.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # create_excel_with_values_2(st.session_state.notes)
            excel_buffer = create_excel_with_values_2(st.session_state.notes)
            st.download_button(
                label="Download Excel",
                data=excel_buffer,
                file_name='price_estimate.xlsx',
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    if st.button("Reset"):
        st.session_state.clear()
        st.success("Reset completed!")



if __name__ == "__main__":
    main()


